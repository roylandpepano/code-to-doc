import os
import docx
import time
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

# Measures the runtime of the program
start = time.time()

# Get the filepath of the folder
mainFolderPath = input("Enter the filepath of the folder: ")

# Get the filename
file_name = os.path.splitext(mainFolderPath)[0]

# Create a document and set the column into 2
document = Document()
section = document.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')

# Open the files using their paths and transfer the content to docx
for path, subdirs, files in os.walk(mainFolderPath):
    for pathFile in files:
        pathFile = os.path.join(path, pathFile)
        fileName = os.path.basename(pathFile)
        srcFile = open(pathFile, "r+", encoding='UTF-8')

        # Add a header [filename.php]
        head = document.add_paragraph()
        head.paragraph_format.space_before = Pt(15)
        head.paragraph_format.space_after = Pt(15)
        head = head.add_run(fileName)
        fhead = head.font
        fhead.name = 'Arial'
        fhead.size = Pt(10)
        fhead.bold = True
        fhead.italic = True
        fhead.underline = True

        # Add the codes
        for line in srcFile:
            line = line.rstrip('\n')
            code = document.add_paragraph()
            code.paragraph_format.space_after = 0
            code = code.add_run(line)
            fcode = code.font
            fcode.name = 'Arial'
            fcode.size = Pt(9)

# Saving to docx file
file_name_docx = file_name + ".docx"
document.save(file_name_docx)

# Saving to pdf for faster viewing
file_name_pdf = file_name + ".pdf"
convert(file_name_docx, file_name_pdf)
print(f"{file_name_docx} and {file_name_pdf} were successfully created!")

end = time.time()
print(f"Program runtime: {end - start} seconds")